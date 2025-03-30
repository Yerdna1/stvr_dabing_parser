import logging
import re
from pathlib import Path
from docling.chunking import HybridChunker
from docling.document_converter import DocumentConverter
import docx
from docx.shared import Inches, RGBColor

_log = logging.getLogger(__name__)

# Column Headers
COLUMN_HEADERS = ["Speaker", "Timecode", "Text", "Scene Marker", "Segment Marker"]

# --- Updated Regular Expression Patterns ---
SPEAKER_BASE = r"[A-ZÁČĎÉÍĹĽŇÓŔŠŤÚÝŽ\s]+\d?"
SPEAKER_PATTERN = rf"({SPEAKER_BASE}):*"
P_MULTIPLE_SPEAKERS = re.compile(rf"^{SPEAKER_PATTERN}(?:,\s*{SPEAKER_PATTERN})+$")
# Pattern to FIND segment marker sequence anywhere
P_SEGMENT_MARKER_FIND = re.compile(r"[-–—]{5,}") # Find 5+ dashes anywhere
P_TIMECODE_FIND = re.compile(r"((?:A\s*)?\b\d{2}:\d{2}(?::\d{2})?(?:-\s*\d{2}:\d{2}(?::\d{2})?)?\b[-–—]*)")
P_SPEAKER_MARKER_FIND = re.compile(rf"^{SPEAKER_PATTERN}\s*(\(.*\))\s*\t")
P_SPEAKER_DASH_FIND = re.compile(rf"^{SPEAKER_PATTERN}\s*[-–—]\s*")
P_SCENE_KEYWORD_FIND = re.compile(r"(INT\.|EXT\.|TITULOK)\s*")
P_PARENS_MARKER_FIND = re.compile(r"(\(.*?\))")

def clean_speaker_name(name: str) -> str:
    """Removes trailing colons (1 or 2) and whitespace from a speaker name."""
    name = name.strip()
    if name.endswith("::"):
        return name[:-2].strip()
    elif name.endswith(":"):
        return name[:-1].strip()
    return name

def parse_document_chunks(chunks: list[str]) -> list[dict[str, str]]:
    """
    Parses lines within chunks using refined element-finding and removal logic.
    Counts segment markers if a line *contains* 5+ dashes.

    Args:
        chunks: A list of text chunks.

    Returns:
        A list of dictionaries, where each dictionary represents a row.
    """
    parsed_rows = []
    segment_marker_count = 0

    for chunk_idx, chunk in enumerate(chunks):
        _log.debug(f"--- Processing Chunk {chunk_idx} ---")
        lines = chunk.splitlines()
        for line_idx, line in enumerate(lines):
            original_line = line.strip()
            if not original_line:
                continue

            # --- Priority 1: Check for full-line Multiple Speakers ---
            multi_speaker_pattern_find = re.compile(rf"({SPEAKER_BASE}):*")
            potential_speakers = multi_speaker_pattern_find.findall(original_line)
            constructed_line = ", ".join(potential_speakers)
            is_multiple_speaker_line = False
            if len(potential_speakers) > 1:
                 cleaned_original = re.sub(r':', '', original_line)
                 cleaned_constructed = re.sub(r':', '', constructed_line)
                 if cleaned_original == cleaned_constructed:
                      is_multiple_speaker_line = True

            if is_multiple_speaker_line:
                 _log.debug(f"Line {line_idx}: Matched as MULTIPLE_SPEAKERS.")
                 speakers = [clean_speaker_name(s) for s in potential_speakers]
                 for speaker in speakers:
                     if speaker:
                         multi_speaker_row = {header: "" for header in COLUMN_HEADERS}
                         multi_speaker_row["Speaker"] = speaker
                         parsed_rows.append(multi_speaker_row)
                 continue # Handled this line

            # --- Priority 2: Find elements within the line ---
            _log.debug(f"Line {line_idx}: Processing line: {repr(original_line)}")
            row_data = {header: "" for header in COLUMN_HEADERS}
            remaining_text = original_line
            found_spans = []

            # 2a: Find Timecode(s)
            timecodes = []
            for match in P_TIMECODE_FIND.finditer(remaining_text):
                timecodes.append(match.group(1))
                found_spans.append(match.span())
            if timecodes:
                row_data["Timecode"] = " ".join(timecodes)
                _log.debug(f"Line {line_idx}: Found Timecode(s): {row_data['Timecode']}")

            # 2b: Find Speaker
            speaker_name = ""
            speaker_match = P_SPEAKER_MARKER_FIND.search(remaining_text)
            if not speaker_match:
                 speaker_match = P_SPEAKER_DASH_FIND.search(remaining_text)

            if speaker_match:
                 speaker_raw = speaker_match.group(1)
                 speaker_name = clean_speaker_name(speaker_raw)
                 row_data["Speaker"] = speaker_name
                 found_spans.append(speaker_match.span(1))
                 if speaker_match.lastindex and speaker_match.lastindex > 1:
                      found_spans.append(speaker_match.span(speaker_match.lastindex))
                 _log.debug(f"Line {line_idx}: Found Speaker: {speaker_name}")

            # 2c: Find Scene Marker(s)
            scene_markers = []
            keyword_match_found = False
            for match in P_SCENE_KEYWORD_FIND.finditer(remaining_text):
                 marker_text = remaining_text[match.start():]
                 scene_markers.append(marker_text.strip())
                 found_spans.append((match.start(), len(remaining_text)))
                 keyword_match_found = True
                 _log.debug(f"Line {line_idx}: Found Scene Keyword Marker: {marker_text.strip()}")
                 break

            for match in P_PARENS_MARKER_FIND.finditer(remaining_text):
                 is_covered = any(span[0] <= match.start() and span[1] >= match.end() for span in found_spans if span != match.span())
                 if not is_covered:
                      scene_markers.append(match.group(1))
                      found_spans.append(match.span())
                      _log.debug(f"Line {line_idx}: Found Parenthetical Marker: {match.group(1)}")

            if scene_markers:
                row_data["Scene Marker"] = " ".join(sorted(list(set(scene_markers)), key=remaining_text.find))

            # 2d: Check for Segment Marker sequence *within* the original line
            if P_SEGMENT_MARKER_FIND.search(original_line):
                segment_marker_count += 1
                row_data["Segment Marker"] = str(segment_marker_count)
                _log.info(f"Line {line_idx}: Found segment marker sequence. Count: {segment_marker_count}. Assigning count to row.")
                # NOTE: We do NOT add the dashes span to found_spans,
                # so they might appear in the Text column if not part of another element.
                # This might be desired or not, depending on exact requirements.

            # 2e: Determine Remaining Text
            processed_text = ""
            last_end = 0
            found_spans.sort()
            for start, end in found_spans:
                if start > last_end:
                    processed_text += remaining_text[last_end:start]
                last_end = max(last_end, end)
            if last_end < len(remaining_text):
                processed_text += remaining_text[last_end:]

            row_data["Text"] = processed_text.strip()
            _log.debug(f"Line {line_idx}: Assigned Remaining Text: {repr(row_data['Text'])}")

            parsed_rows.append(row_data)

    return parsed_rows


def export_to_docx_table(parsed_rows: list[dict[str, str]], output_path: Path):
    """Exports the parsed elements to a DOCX file in a table format with speaker coloring."""
    _log.info(f"Exporting {len(parsed_rows)} rows to table in {output_path}...")
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=len(COLUMN_HEADERS))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(COLUMN_HEADERS):
        hdr_cells[i].text = header

    for row_data in parsed_rows:
        row_cells = table.add_row().cells
        for i, header in enumerate(COLUMN_HEADERS):
            content = row_data.get(header, "").replace('\t', '    ')
            cell = row_cells[i]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.TOP
            run = paragraph.add_run(content)
            if header == "Speaker" and content:
                run.font.color.rgb = RGBColor(255, 0, 0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    _log.info("Export complete.")


def parse_and_export(source_path: Path, output_docx_path: Path):
    """Main function to convert, chunk, parse, and export."""
    _log.info(f"Starting processing for {source_path}...")
    # 1. Convert document
    _log.info(f"Loading and converting document from {source_path}...")
    try:
        converter = DocumentConverter()
        conv_result = converter.convert(source=source_path)
        if not conv_result or not conv_result.document:
             _log.error(f"Failed to convert document from {source_path}")
             return
        doc = conv_result.document
    except FileNotFoundError:
        _log.error(f"Error: Source file not found at {source_path}")
        return
    except Exception as e:
        _log.error(f"Error during document conversion: {e}")
        return

    # 2. Chunk document
    _log.info("Initializing HybridChunker...")
    chunker = HybridChunker()
    _log.info("Chunking document...")
    try:
        chunk_iter = chunker.chunk(dl_doc=doc)
        serialized_chunks = [chunker.serialize(chunk=chunk) for chunk in chunk_iter]
    except Exception as e:
        _log.error(f"Error during chunking: {e}")
        return

    _log.info(f"Generated {len(serialized_chunks)} chunks.")
    if not serialized_chunks:
        _log.warning("No chunks were generated.")
        return

    # 3. Parse chunks using the revised logic
    _log.info("Parsing generated chunks using element-finding logic...")
    parsed_row_data = parse_document_chunks(serialized_chunks)
    if not parsed_row_data:
        _log.warning("Parsing did not yield any elements.")
        return

    # 4. Export to DOCX Table
    export_to_docx_table(parsed_row_data, output_docx_path)

    _log.info("Processing finished.")


if __name__ == "__main__":
    # Keep logging level at DEBUG for verification
    logging.basicConfig(level=logging.DEBUG, format='%(levelname)s:%(name)s:%(message)s')
    input_source_path = Path("test4.docx")
    output_docx_file = Path("output/parsed_output.docx")

    if not input_source_path.exists():
        _log.error(f"Input file {input_source_path} does not exist.")
    else:
        parse_and_export(input_source_path, output_docx_file)
