"""
DOCX Export Agent for screenplay formatting and export
"""
import os
import re
import sys
import traceback
from typing import Dict, List, Any, Optional
import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class DocxExportAgent:
    """Agent for exporting screenplay data to formatted DOCX files."""
    
    def __init__(self):
        """Initialize the DOCX export agent."""
        # Define colors for different elements
        self.RED_COLOR = RGBColor(255, 0, 0)  # Speaker names
        self.BLUE_COLOR = RGBColor(0, 0, 255)  # Scene headers (INT/EXT)
        self.BLACK_COLOR = RGBColor(0, 0, 0)   # Regular text
        
        # Define font and paragraph settings
        self.FONT_NAME = 'Verdana'
        self.FONT_SIZE = 13  # Points
        self.LINE_SPACING = 1.5
        
    def export_to_docx(self, segments: List[Dict], output_path: str, episode_number: Optional[str] = None) -> str:
        """
        Export screenplay segments to a formatted DOCX file.
        
        Args:
            segments: List of screenplay segments
            output_path: Path to save the DOCX file
            episode_number: Optional episode number for segment numbering
            
        Returns:
            Path to the created DOCX file
        """
        try:
            # Log input parameters
            st.write(f"🔍 Starting DOCX export process")
            st.write(f"📁 Original output path: {output_path}")
            st.write(f"🔢 Episode number: {episode_number}")
            st.write(f"📊 Number of segments to process: {len(segments)}")
            
            # Fix path for Windows - ensure directory exists
            output_path = os.path.normpath(output_path)
            st.write(f"📁 Normalized output path: {output_path}")
            
            # Check if directory exists
            dir_path = os.path.dirname(output_path)
            st.write(f"📂 Directory path: {dir_path}")
            
            if not os.path.exists(dir_path):
                st.write(f"📂 Creating directory: {dir_path}")
                os.makedirs(dir_path, exist_ok=True)
            else:
                st.write(f"📂 Directory already exists")
            
            # Log current working directory
            st.write(f"📂 Current working directory: {os.getcwd()}")
            
            # Check write permissions
            try:
                if dir_path:
                    test_file = os.path.join(dir_path, 'test_write.tmp')
                else:
                    test_file = 'test_write.tmp'
                with open(test_file, 'w') as f:
                    f.write('test')
                os.remove(test_file)
                st.write(f"✅ Write permission test passed")
            except Exception as e:
                st.error(f"❌ Write permission test failed: {str(e)}")
            
            # Create a new document
            st.write(f"📝 Creating new Document")
            doc = Document()
            
            # Configure document settings
            st.write(f"⚙️ Configuring document settings")
            style = doc.styles['Normal']
            style.font.name = self.FONT_NAME
            style.font.size = Pt(self.FONT_SIZE)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = self.LINE_SPACING
            
            # Set margins
            section = doc.sections[0]
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            # Create a table for the content (2 columns: speaker and content)
            st.write(f"📊 Creating content table")
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths
            table.columns[0].width = Cm(5)    # Speaker column
            table.columns[1].width = Cm(14)   # Content column
            
            # Make table borders invisible - simplified approach
            st.write(f"🔲 Setting invisible borders")
            # We'll handle this more simply for compatibility
            
            # Calculate available width for separators
            available_width = self._get_available_width(section)
            separator_length = available_width - 16  # Leave some margin
            st.write(f"📏 Separator length: {separator_length} characters")
            
            # Track the current segment number
            segment_count = 0
            
            # Process each segment
            st.write(f"🔄 Processing segments")
            processed_count = 0
            
            for segment in segments:
                try:
                    # Check if this is a segment marker
                    is_segment_marker = (segment.get("type") == "segment_marker" or 
                                    ("timecode" in segment and self._is_segment_marker(segment.get("timecode", ""))))
                    
                    if is_segment_marker:
                        # Increment segment count if not already provided
                        if "segment_number" in segment:
                            segment_count = segment["segment_number"]
                        else:
                            segment_count += 1
                        
                        # Format segment number
                        if episode_number:
                            segment_number = f"{int(episode_number)}{segment_count:02d}"
                        else:
                            segment_number = f"{segment_count:02d}"
                        
                        # Add separator line (dashes)
                        self._add_separator_row(table, '-' * separator_length)
                        
                        # Add segment number line
                        timecode = segment.get("timecode", "").split('-')[0].strip('*')  # Extract the timecode part
                        
                        # Calculate spacing to position segment number near the right margin
                        right_margin = 2  # Number of spaces from right edge
                        spacing = separator_length - len(timecode) - len(segment_number) - right_margin
                        
                        # Format the segment text with the timecode at left, segment number at right
                        segment_text = f"{timecode}" + " " * spacing + f"{segment_number}"
                        
                        # Add to table with bold text but no spacing after
                        self._add_segment_number_row(table, segment_text)
                        
                    # Handle scene headers (INT/EXT)
                    elif segment.get("text", "").upper().startswith("INT") or segment.get("text", "").upper().startswith("EXT"):
                        scene_text = segment.get("text", "")
                        
                        # Try to extract the scene type (INT/EXT)
                        scene_match = re.match(r'^(INT|EXT)\.?\s*(.*)$', scene_text, re.IGNORECASE)
                        if scene_match:
                            scene_type = scene_match.group(1).upper()
                            scene_desc = scene_match.group(2).strip()
                            self._add_split_content(table, scene_type, scene_desc, speaker_color=self.BLUE_COLOR)
                        else:
                            self._add_split_content(table, "", scene_text)
                        
                    # Handle speakers with dialogue
                    elif "speaker" in segment:
                        speaker = segment.get("speaker", "")
                        text = segment.get("text", "")
                        
                        # If there's a timecode, add it before the speaker
                        if "timecode" in segment and segment["timecode"] and not self._is_segment_marker(segment["timecode"]):
                            speaker = f"{segment['timecode']} {speaker}"
                        
                        self._add_split_content(table, speaker, text, speaker_color=self.RED_COLOR)
                    
                    # Handle any other content
                    else:
                        self._add_split_content(table, "", segment.get("text", ""))
                    
                    processed_count += 1
                    
                    # Log progress occasionally
                    if processed_count % 50 == 0:
                        st.write(f"🔄 Processed {processed_count}/{len(segments)} segments")
                        
                except Exception as e:
                    st.error(f"❌ Error processing segment {processed_count}: {str(e)}")
                    st.error(f"Segment data: {segment}")
            
            st.write(f"✅ Processed all {processed_count} segments")
            
            # Save the document
            st.write(f"💾 Attempting to save document to: {output_path}")
            
            try:
                doc.save(output_path)
                st.success(f"✅ Document successfully saved to: {output_path}")
                return output_path
            except Exception as e:
                error_msg = f"❌ Error during save operation: {str(e)}"
                st.error(error_msg)
                # Get detailed error information
                exc_type, exc_value, exc_traceback = sys.exc_info()
                error_details = traceback.format_exception(exc_type, exc_value, exc_traceback)
                st.write("🔍 Detailed error trace:")
                for line in error_details:
                    st.write(line)
                return ""
                
        except Exception as e:
            error_msg = f"❌ Error in DOCX export process: {str(e)}"
            st.error(error_msg)
            # Get detailed error information
            exc_type, exc_value, exc_traceback = sys.exc_info()
            error_details = traceback.format_exception(exc_type, exc_value, exc_traceback)
            st.write("🔍 Detailed error trace:")
            for line in error_details:
                st.write(line)
            return ""
    
    def _is_segment_marker(self, timecode: str) -> bool:
        """Check if a timecode is a segment marker (contains multiple dashes)."""
        return bool(re.search(r'\*?\*?[\d:]+[-]{5,}', timecode) or 
                   (timecode and len(timecode.strip()) >= 4 and '-' in timecode and ':' in timecode))
    
    def _get_available_width(self, section) -> int:
        """Calculate available width in characters based on page settings."""
        # Get page width in inches
        page_width = section.page_width.inches
        # Subtract margins
        available_width = page_width - section.left_margin.inches - section.right_margin.inches + 2.7
        # Convert to approximate character count (assuming each char is ~0.1 inches)
        char_count = int(available_width / 0.1)
        return char_count
    
    def _add_separator_row(self, table, text, is_bold=False):
        """Add a separator row with dashes."""
        try:
            row = table.add_row()
            cell = row.cells[0].merge(row.cells[1])
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            
            # Apply formatting
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = is_bold
            
            # Set spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)  # Add space after separator
            paragraph.paragraph_format.line_spacing = self.LINE_SPACING
            
            return row
        except Exception as e:
            st.error(f"❌ Error in _add_separator_row: {str(e)}")
            raise
    
    def _add_segment_number_row(self, table, text):
        """Add a segment number row with proper formatting."""
        try:
            row = table.add_row()
            cell = row.cells[0].merge(row.cells[1])
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            
            # Apply formatting
            run.font.name = self.FONT_NAME
            run.font.size = Pt(self.FONT_SIZE)
            run.font.bold = True  # Bold for segment numbers
            
            # Set spacing - NO spacing after segment number
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)  # No space after segment number
            paragraph.paragraph_format.line_spacing = self.LINE_SPACING
            
            return row
        except Exception as e:
            st.error(f"❌ Error in _add_segment_number_row: {str(e)}")
            raise
    
    def _add_split_content(self, table, speaker, content, speaker_color=None):
        """Add a row with speaker and content in separate columns."""
        try:
            row = table.add_row()
            
            # Speaker cell
            if speaker:
                speaker_para = row.cells[0].paragraphs[0]
                speaker_run = speaker_para.add_run(speaker)
                
                # Apply formatting
                speaker_run.font.name = self.FONT_NAME
                speaker_run.font.size = Pt(self.FONT_SIZE)
                if speaker_color:
                    speaker_run.font.color.rgb = speaker_color
                
                # Set spacing
                speaker_para.paragraph_format.space_before = Pt(0)
                speaker_para.paragraph_format.space_after = Pt(0)
                speaker_para.paragraph_format.line_spacing = self.LINE_SPACING
            
            # Content cell
            if content:
                content_para = row.cells[1].paragraphs[0]
                content_run = content_para.add_run(content)
                
                # Apply formatting
                content_run.font.name = self.FONT_NAME
                content_run.font.size = Pt(self.FONT_SIZE)
                
                # Set spacing
                content_para.paragraph_format.space_before = Pt(0)
                content_para.paragraph_format.space_after = Pt(0)
                content_para.paragraph_format.line_spacing = self.LINE_SPACING
            
            return row
        except Exception as e:
            st.error(f"❌ Error in _add_split_content: {str(e)}")
            st.error(f"Speaker: {speaker}, Content: {content[:30]}...")
            raise